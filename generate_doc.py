"""
WebGuard AI — College Project Document Generator
Run: python generate_doc.py
Output: WebGuard_AI_Project_Document.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

# ── helpers ──────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=(189, 178, 255)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code, lang=""):
    """Monospace grey box for code."""
    style = doc.styles['Normal']
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def hr(doc):
    doc.add_paragraph("─" * 80)

# ── Read main source file ─────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
APP_PY = os.path.join(SCRIPT_DIR, "ai_service", "app.py")
with open(APP_PY, encoding="utf-8") as f:
    app_py_code = f.read()

# ── Build document ────────────────────────────────────────────────────────────
doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("WebGuard AI", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size  = Pt(28)
    run.font.color.rgb = RGBColor(189, 178, 255)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Intelligent Log Analysis & Threat Detection Platform")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 220, 160)

doc.add_paragraph()
meta = [
    ("Project Type",  "BTech Final Year Project"),
    ("Technology",    "Python · Flask · Scikit-Learn · React · Node.js · MongoDB"),
    ("Detection",     "Isolation Forest ML + Deterministic Rule Engine"),
    ("Date",          datetime.date.today().strftime("%B %Y")),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(val);          r2.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Project Overview")
add_para(doc,
    "WebGuard AI is an enterprise-grade cybersecurity platform that analyses Apache/Nginx "
    "web server access logs to detect threats using a two-tier approach: a deterministic "
    "rule-based engine (52+ signatures across 8 attack categories) combined with an "
    "Isolation Forest machine learning model. The system delivers AI-generated expert "
    "security briefings and implements a full role-based analyst workflow.")

doc.add_paragraph()
add_heading(doc, "1.1 Architecture", level=2)
arch_rows = [
    ("Layer",          "Component",            "Technology"),
    ("ML Microservice","Threat detection API",  "Python · Flask · Scikit-Learn · Polars"),
    ("Backend API",    "Log processing server", "Node.js · Express · MongoDB"),
    ("Frontend",       "SOC dashboard",         "React · TypeScript · Vite · Tailwind CSS"),
    ("Database",       "Persistence",           "MongoDB Atlas / local"),
]
tbl = doc.add_table(rows=len(arch_rows), cols=3)
tbl.style = 'Table Grid'
for i, row_data in enumerate(arch_rows):
    cells = tbl.rows[i].cells
    for j, val in enumerate(row_data):
        cells[j].text = val
        run = cells[j].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if i == 0:
            run.bold = True

doc.add_paragraph()
add_heading(doc, "1.2 Key Features", level=2)
features = [
    "52+ deterministic attack signatures (SQLi, XSS, Path Traversal, RCE, Brute-Force, Recon, Exfil, Scanners)",
    "Isolation Forest anomaly model with RobustScaler preprocessing and 13 engineered features",
    "IP-level behavioural intelligence: rapid burst detection, enumeration, error-rate analysis",
    "Local expert AI briefing engine — generates professional security reports without any external API",
    "Role-based access: Log Contributor → Security Analyst → Security Manager escalation workflow",
    "Live threat stats synced to dashboard header after each analysis run",
    "Multi-format download: Expert CSV, Executive Briefing TXT, Raw JSON",
]
for f in features:
    add_bullet(doc, f)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. SOFTWARE & DEPENDENCIES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Software & Dependencies")

add_heading(doc, "2.1 Python (ai_service/)", level=2)
py_deps = [
    ("Flask 3.x",        "HTTP microservice framework"),
    ("flask-cors",       "Cross-Origin Resource Sharing"),
    ("scikit-learn 1.x", "Isolation Forest anomaly detection model"),
    ("polars",           "Fast DataFrame for IP profiling"),
    ("numpy",            "Numerical feature engineering"),
    ("joblib",           "Model serialisation / deserialisation"),
]
tbl2 = doc.add_table(rows=len(py_deps)+1, cols=2)
tbl2.style = 'Table Grid'
tbl2.rows[0].cells[0].text = "Package"; tbl2.rows[0].cells[1].text = "Purpose"
for r in tbl2.rows[0].cells:
    r.paragraphs[0].runs[0].bold = True; r.paragraphs[0].runs[0].font.size = Pt(9)
for i, (pkg, desc) in enumerate(py_deps, 1):
    tbl2.rows[i].cells[0].text = pkg; tbl2.rows[i].cells[1].text = desc
    for c in tbl2.rows[i].cells:
        c.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
add_heading(doc, "2.2 Node.js (server/)", level=2)
node_deps = [
    ("Express 5",     "REST API server"),
    ("multer",        "Multipart log file uploads"),
    ("mongoose",      "MongoDB ODM"),
    ("node-fetch",    "HTTP calls to Python microservice"),
    ("dotenv",        "Environment variable management"),
    ("cors",          "CORS middleware"),
]
tbl3 = doc.add_table(rows=len(node_deps)+1, cols=2)
tbl3.style = 'Table Grid'
tbl3.rows[0].cells[0].text = "Package"; tbl3.rows[0].cells[1].text = "Purpose"
for r in tbl3.rows[0].cells:
    r.paragraphs[0].runs[0].bold = True; r.paragraphs[0].runs[0].font.size = Pt(9)
for i, (pkg, desc) in enumerate(node_deps, 1):
    tbl3.rows[i].cells[0].text = pkg; tbl3.rows[i].cells[1].text = desc
    for c in tbl3.rows[i].cells:
        c.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
add_heading(doc, "2.3 Frontend", level=2)
fe_deps = [
    ("React 19 + TypeScript", "UI framework"),
    ("Vite 7",                "Build tool and dev server"),
    ("Tailwind CSS 3",        "Utility-first CSS framework"),
    ("Framer Motion",         "Animation library"),
    ("Recharts",              "Threat radar and charts"),
    ("Lucide React",          "Icon library"),
    ("React Router DOM 7",    "Client-side routing"),
]
tbl4 = doc.add_table(rows=len(fe_deps)+1, cols=2)
tbl4.style = 'Table Grid'
tbl4.rows[0].cells[0].text = "Package"; tbl4.rows[0].cells[1].text = "Purpose"
for r in tbl4.rows[0].cells:
    r.paragraphs[0].runs[0].bold = True; r.paragraphs[0].runs[0].font.size = Pt(9)
for i, (pkg, desc) in enumerate(fe_deps, 1):
    tbl4.rows[i].cells[0].text = pkg; tbl4.rows[i].cells[1].text = desc
    for c in tbl4.rows[i].cells:
        c.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. README — SETUP & EXECUTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Setup & Execution (README)")

add_heading(doc, "3.1 Prerequisites", level=2)
prereqs = [
    "Python 3.11+ (with pip)",
    "Node.js 20 LTS (with npm)",
    "MongoDB 7 running locally (or MongoDB Atlas connection string)",
    "Git",
]
for p in prereqs:
    add_bullet(doc, p)

doc.add_paragraph()
add_heading(doc, "3.2 Installation", level=2)

add_para(doc, "Step 1 — Clone and install frontend + backend dependencies:", bold=True, size=10)
add_code_block(doc,
"git clone <repo-url> webguard-ai\n"
"cd webguard-ai\n"
"npm install")

add_para(doc, "Step 2 — Install Python ML service dependencies:", bold=True, size=10)
add_code_block(doc,
"cd ai_service\n"
"pip install flask flask-cors scikit-learn polars numpy joblib")

add_para(doc, "Step 3 — Configure environment variables:", bold=True, size=10)
add_code_block(doc,
"# Create webguard-ai/server/.env\n"
"MONGO_URI=mongodb://localhost:27017/webguard\n"
"PORT=5000\n"
"# Optional — enables Mistral LLM briefings (not required)\n"
"# HUGGINGFACE_API_KEY=your_key_here")

doc.add_paragraph()
add_heading(doc, "3.3 Running the Application", level=2)
add_para(doc, "Start all three services in separate terminals:", bold=False, size=10)

add_para(doc, "Terminal 1 — Python ML Microservice (port 5001):", bold=True, size=10)
add_code_block(doc, "cd webguard-ai/ai_service\npython app.py")

add_para(doc, "Terminal 2 — Node.js Backend API (port 5000):", bold=True, size=10)
add_code_block(doc, "cd webguard-ai\nnode server/index.js")

add_para(doc, "Terminal 3 — Vite Frontend Dev Server (port 5173):", bold=True, size=10)
add_code_block(doc, "cd webguard-ai\nnpm run dev")

add_para(doc, "Open http://localhost:5173 in your browser.", bold=False, size=10)

doc.add_paragraph()
add_heading(doc, "3.4 Demo Login Credentials", level=2)
creds = [
    ("Log Contributor",  "contributor@webguard.ai", "contributor123"),
    ("Security Analyst", "analyst@webguard.ai",     "analyst123"),
    ("Security Manager", "manager@webguard.ai",     "manager123"),
]
tbl5 = doc.add_table(rows=len(creds)+1, cols=3)
tbl5.style = 'Table Grid'
for j, h in enumerate(["Role", "Email", "Password"]):
    tbl5.rows[0].cells[j].text = h
    tbl5.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    tbl5.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(9)
for i, (role, email, pwd) in enumerate(creds, 1):
    for j, val in enumerate([role, email, pwd]):
        tbl5.rows[i].cells[j].text = val
        tbl5.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()
add_heading(doc, "3.5 Testing with Sample Log File", level=2)
steps = [
    "Log in as Log Contributor → Upload test_logs.txt",
    "Log in as Security Analyst → Click ⚡ on the queued file",
    "Review flagged threats in the analysis table (red rows = anomalies)",
    "Click 'Generate AI-Augmented Report' to produce an expert security briefing",
    "Log in as Security Manager → View the briefing and download the report",
]
for s in steps:
    add_numbered(doc, s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MAIN SOURCE FILE — ai_service/app.py
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Main Source File — ai_service/app.py")
add_para(doc,
    "This is the core ML microservice. It implements log parsing, the deterministic rule "
    "engine (8 attack categories, 52+ regex signatures), IP-level behavioural intelligence, "
    "ML feature engineering, and the Isolation Forest inference pipeline.",
    size=10)

doc.add_paragraph()

# Split code into sections to avoid one massive block
sections_code = [
    ("Imports & Model Loading", 1, 22),
    ("Log Parsing (Regex)", 23, 59),
    ("Threat Rules — SQLi, Traversal, Recon, Attack Tools", 60, 143),
    ("Threat Rules — XSS, Command Injection, Brute Force, Exfil", 144, 212),
    ("IP Behavioural Intelligence", 213, 238),
    ("Rule-Based Detection Engine", 239, 309),
    ("ML Feature Engineering", 310, 368),
    ("Main Analysis Pipeline", 369, 436),
    ("ML Explanation Helper & Flask API", 437, 496),
]

lines = app_py_code.splitlines()
for section_title, start, end in sections_code:
    add_heading(doc, f"4.{sections_code.index((section_title,start,end))+1}  {section_title}", level=2)
    chunk = "\n".join(lines[start-1:end])
    add_code_block(doc, chunk)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SYSTEM DESIGN NOTES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. System Design Notes")

add_heading(doc, "5.1 Detection Strategy — Defence in Depth", level=2)
add_para(doc,
    "WebGuard AI uses a two-tier detection approach. The deterministic rule engine fires "
    "first and always overrides the ML verdict for known attack signatures. The Isolation "
    "Forest then catches novel or statistically unusual patterns not covered by rules. "
    "This ensures zero false negatives for known attacks while still flagging unknown threats.")

add_bullet(doc, "Rules fire on: URL content, User-Agent string, HTTP status code, auth failure count")
add_bullet(doc, "ML fires on:  13 engineered features including hour-of-day, url_length, ip_error_rate, ua_is_script")
add_bullet(doc, "Combined:  is_anomaly = rule_threat OR ml_anomaly")

doc.add_paragraph()
add_heading(doc, "5.2 Key Fix — URL Regex", level=2)
add_para(doc,
    "The original system used \\S+ to capture URLs, which truncated any URL containing "
    "spaces — causing all SQL injection payloads (e.g. UNION SELECT username,password) "
    "to be silently cut before analysis. The fix uses a lazy .+? anchored by \\s+HTTP/, "
    "ensuring the full payload is captured:")
add_code_block(doc,
"# BEFORE (broken — truncates at first space)\n"
"r'(?P<url>\\S+)'\n\n"
"# AFTER (correct — captures full URL including spaces in payloads)\n"
"r'(?P<url>.+?)\\s+(?P<protocol>HTTP/[\\d.]+)'")

doc.add_paragraph()
add_heading(doc, "5.3 API Transport Fix — CSV → JSON", level=2)
add_para(doc,
    "The Node.js /process endpoint originally returned CSV. SQL injection URLs contain "
    "commas (e.g. UNION SELECT username,password), causing column misalignment in the CSV "
    "parser and silent data loss. The fix changed the transport to JSON, where fields are "
    "read by name, eliminating all parsing fragility.")

doc.add_paragraph()
add_heading(doc, "5.4 AI Report Engine", level=2)
add_para(doc,
    "Rather than requiring an external HuggingFace API key, WebGuard AI includes a local "
    "expert analysis engine (server/services/aiReporting.js) that:")
add_bullet(doc, "Categorises each threat by attack type")
add_bullet(doc, "Profiles top attacking IPs by frequency")
add_bullet(doc, "Computes risk level (LOW/MEDIUM/HIGH/CRITICAL) from anomaly rate and severity")
add_bullet(doc, "Selects category-specific remediation playbooks (e.g. SQLi → WAF + parameterised queries)")
add_bullet(doc, "Falls back to HuggingFace Mistral-7B if HUGGINGFACE_API_KEY is present")

doc.add_paragraph()
add_heading(doc, "5.5 Project File Structure", level=2)
add_code_block(doc,
"webguard-ai/\n"
"├── ai_service/           # Python ML microservice (port 5001)\n"
"│   ├── app.py            # ← MAIN FILE (this document)\n"
"│   ├── isolation_forest_model.pkl\n"
"│   ├── robust_scaler.pkl\n"
"│   └── ip_profile_train.pkl\n"
"├── server/               # Node.js backend (port 5000)\n"
"│   ├── index.js\n"
"│   ├── routes/\n"
"│   │   ├── analysis.js   # Log processing + report endpoints\n"
"│   │   └── reports.js\n"
"│   ├── models/           # MongoDB schemas\n"
"│   └── services/\n"
"│       └── aiReporting.js # Expert briefing engine\n"
"├── src/                  # React frontend (port 5173)\n"
"│   ├── pages/\n"
"│   │   ├── Landing.tsx\n"
"│   │   └── Login.tsx\n"
"│   ├── components/\n"
"│   │   ├── dashboard/    # AnalyzeView, ManagerView, CommandView\n"
"│   │   └── layout/       # TopBar, Sidebar, AppLayout\n"
"│   └── context/\n"
"│       ├── AuthContext.tsx\n"
"│       └── SecurityContext.tsx\n"
"└── package.json")

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(SCRIPT_DIR), "WebGuard_AI_Project_Document.docx")
out_path = os.path.join(SCRIPT_DIR, "WebGuard_AI_Project_Document.docx")
doc.save(out_path)
print(f"✅  Document saved: {out_path}")
print(f"    Pages (approx): ~10-11")
